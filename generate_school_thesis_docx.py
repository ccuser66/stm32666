# -*- coding: gbk -*-
from __future__ import annotations

import re
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


ROOT = Path(r"e:\designdemo\STM32CodeDemo")
TEMPLATE_PATH = ROOT / "论文参考" / "毕业论文-代云吉-终稿.docx"
SOURCE_TEXT_PATH = ROOT / "毕业论文初稿_基于STM32的音频信号采集与可视化系统设计与实现.txt"
OUTPUT_DOCX_PATH = ROOT / "基于STM32的音频信号采集与可视化系统设计与实现_学校模板版.docx"


def read_text_auto(path: Path) -> str:
    for encoding in ("utf-8", "gbk", "utf-8-sig"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    raise UnicodeDecodeError("auto", b"", 0, 1, f"无法解码文件: {path}")


def set_paragraph_text(paragraph, text: str) -> None:
    element = paragraph._element
    ppr = qn("w:pPr")
    for child in list(element):
        if child.tag != ppr:
            element.remove(child)
    if text:
        paragraph.add_run(text)


def split_zh_paragraph(text: str, parts: int = 2) -> list[str]:
    text = text.strip()
    if not text:
        return [""] * parts
    sentences = [segment for segment in re.split(r"(?<=[。！？])", text) if segment]
    if len(sentences) < parts:
        return [text] + [""] * (parts - 1)
    chunk_size = max(1, len(sentences) // parts)
    result = []
    cursor = 0
    for index in range(parts - 1):
        next_cursor = min(len(sentences), cursor + chunk_size)
        result.append("".join(sentences[cursor:next_cursor]).strip())
        cursor = next_cursor
    result.append("".join(sentences[cursor:]).strip())
    while len(result) < parts:
        result.append("")
    return result[:parts]


def split_en_paragraph(text: str, parts: int = 2) -> list[str]:
    text = text.strip()
    if not text:
        return [""] * parts
    sentences = [segment for segment in re.split(r"(?<=[.!?])\s+", text) if segment]
    if len(sentences) < parts:
        return [text] + [""] * (parts - 1)
    chunk_size = max(1, len(sentences) // parts)
    result = []
    cursor = 0
    for index in range(parts - 1):
        next_cursor = min(len(sentences), cursor + chunk_size)
        result.append(" ".join(sentences[cursor:next_cursor]).strip())
        cursor = next_cursor
    result.append(" ".join(sentences[cursor:]).strip())
    while len(result) < parts:
        result.append("")
    return result[:parts]


def extract_blocks(lines: list[str], start: int, end: int) -> list[str]:
    blocks: list[str] = []
    current: list[str] = []
    for raw_line in lines[start:end]:
        line = raw_line.strip()
        if not line:
            if current:
                blocks.append("".join(current).strip())
                current = []
            continue
        current.append(line)
    if current:
        blocks.append("".join(current).strip())
    return blocks


def is_heading_1(line: str) -> bool:
    return bool(
        re.match(r"^第\d+章", line)
        or line in {"致谢", "参考文献"}
        or re.match(r"^附录[A-Z]", line)
    )


def is_heading_2(line: str) -> bool:
    return bool(re.match(r"^\d+\.\d+(?!\.)", line))


def is_heading_3(line: str) -> bool:
    return bool(re.match(r"^\d+\.\d+\.\d+", line))


def is_reference_line(line: str) -> bool:
    return bool(re.match(r"^\[\d+\]", line))


def is_caption_line(line: str) -> bool:
    return bool(re.match(r"^[图表]\d+-\d+", line))


def is_list_item(line: str) -> bool:
    return bool(re.match(r"^\d+\.\s", line))


def is_protocol_line(line: str) -> bool:
    command_prefixes = (
        "STAT:",
        "WAV:",
        "FFT:",
        "VU:",
        "GET",
        "PING",
        "MODE:",
        "WIN:",
        "TH:",
        "DENOISE:",
        "CTRL:",
        "VIEW:",
        "SRC:",
    )
    return line.startswith(command_prefixes)


def build_body_items(lines: list[str], start: int) -> list[str]:
    items: list[str] = []
    current: list[str] = []

    def flush_current() -> None:
        nonlocal current
        if current:
            items.append("".join(current).strip())
            current = []

    for raw_line in lines[start:]:
        line = raw_line.strip()
        if not line:
            flush_current()
            continue

        standalone = (
            is_heading_1(line)
            or is_heading_2(line)
            or is_heading_3(line)
            or is_reference_line(line)
            or is_caption_line(line)
            or is_list_item(line)
            or is_protocol_line(line)
        )

        if standalone:
            flush_current()
            items.append(line)
        else:
            current.append(line)

    flush_current()
    return items


def remove_body_from_first_heading(document: Document) -> None:
    body_start = None
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith("第1章") and paragraph.style.name == "Heading 1":
            body_start = paragraph._element
            break

    if body_start is None:
        raise RuntimeError("模板中未找到正文起始标题。")

    body = document._body._element
    children = list(body)
    start_index = children.index(body_start)
    sectpr_tag = qn("w:sectPr")
    for child in children[start_index:]:
        if child.tag != sectpr_tag:
            body.remove(child)


def add_body_content(document: Document, items: list[str]) -> None:
    reference_style_name = "参考文献正文" if any(style.name == "参考文献正文" for style in document.styles) else "Normal"
    heading_1_seen = False

    for item in items:
        if is_heading_1(item):
            if heading_1_seen:
                document.add_page_break()
            paragraph = document.add_paragraph(item, style="Heading 1")
            heading_1_seen = True
        elif is_heading_3(item):
            paragraph = document.add_paragraph(item, style="Heading 3")
        elif is_heading_2(item):
            paragraph = document.add_paragraph(item, style="Heading 2")
        elif is_reference_line(item):
            paragraph = document.add_paragraph(item, style=reference_style_name)
        else:
            paragraph = document.add_paragraph(item, style="Normal")

        if is_caption_line(item):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def update_word_fields(docx_path: Path) -> None:
    powershell_script = rf"""
$path = (Resolve-Path '{docx_path}').Path
$word = $null
$doc = $null
try {{
    $word = New-Object -ComObject Word.Application
    $word.Visible = $false
    $doc = $word.Documents.Open($path)
    $doc.Fields.Update() | Out-Null
    foreach ($toc in $doc.TablesOfContents) {{ $toc.Update() | Out-Null }}
    $doc.Save()
}}
finally {{
    if ($doc -ne $null) {{ try {{ $doc.Close() }} catch {{}} }}
    if ($word -ne $null) {{ try {{ $word.Quit() }} catch {{}} }}
}}
""".strip()
    result = subprocess.run(
        ["powershell.exe", "-NoProfile", "-Command", powershell_script],
        check=False,
        capture_output=True,
        text=True,
    )
    if result.returncode != 0:
        stderr_text = (result.stderr or "").strip()
        stdout_text = (result.stdout or "").strip()
        message = stderr_text or stdout_text or "未知错误"
        print(f"Warning: Word 字段自动刷新失败，文档已生成，可在 Word 中手动更新目录。原因: {message}")


def main() -> None:
    text = read_text_auto(SOURCE_TEXT_PATH)
    lines = text.splitlines()

    zh_abs_title_idx = next(i for i, line in enumerate(lines) if line.strip() == "摘    要")
    zh_kw_idx = next(i for i, line in enumerate(lines) if line.strip().startswith("关键词："))
    en_abs_title_idx = next(i for i, line in enumerate(lines) if line.strip() == "ABSTRACT")
    en_kw_idx = next(i for i, line in enumerate(lines) if line.strip().startswith("Key words:"))
    body_idx = next(i for i, line in enumerate(lines) if line.strip() == "第1章 绪论")

    zh_blocks = extract_blocks(lines, zh_abs_title_idx + 1, zh_kw_idx)
    en_blocks = extract_blocks(lines, en_abs_title_idx + 1, en_kw_idx)

    zh_abstract = "".join(zh_blocks).strip()
    zh_keywords = lines[zh_kw_idx].strip()
    en_abstract = "".join(en_blocks).strip()
    en_keywords = lines[en_kw_idx].strip()
    body_items = build_body_items(lines, body_idx)

    zh_parts = split_zh_paragraph(zh_abstract, parts=2)
    en_parts = split_en_paragraph(en_abstract, parts=2)

    document = Document(str(TEMPLATE_PATH))
    nonempty = [paragraph for paragraph in document.paragraphs if paragraph.text.strip()]

    cover_map = {
        0: "本科毕业设计（论文）",
        1: "题目： 基于STM32的音频信号采集与可视化系统设计与实现",
        2: "学       院：         ____________",
        3: "专       业：                 ____________",
        4: "学 生 姓 名：                 ____________",
        5: "学       号：       ____________",
        6: "指 导 教 师：           ____________",
        7: "评 阅 教 师：          ____________",
        8: "完 成 时 间：      2026年4月25日",
    }
    for index, value in cover_map.items():
        set_paragraph_text(nonempty[index], value)

    declaration_map = {
        14: "作者签名（亲笔）：                      年   月   日",
        18: "作者签名（亲笔）：                      年   月   日",
        19: "导师签名（亲笔）：                      年   月   日",
    }
    for index, value in declaration_map.items():
        set_paragraph_text(nonempty[index], value)

    set_paragraph_text(nonempty[20], "摘    要")
    set_paragraph_text(nonempty[21], zh_parts[0])
    set_paragraph_text(nonempty[22], zh_parts[1])
    set_paragraph_text(nonempty[23], zh_keywords)

    set_paragraph_text(nonempty[25], "Abstract")
    set_paragraph_text(nonempty[26], en_parts[0])
    set_paragraph_text(nonempty[27], en_parts[1])
    set_paragraph_text(nonempty[28], en_keywords)

    remove_body_from_first_heading(document)
    add_body_content(document, body_items)
    document.save(str(OUTPUT_DOCX_PATH))
    update_word_fields(OUTPUT_DOCX_PATH)
    print(f"Generated: {OUTPUT_DOCX_PATH}")


if __name__ == "__main__":
    main()